from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('GLOSARIO DE TÉRMINOS TÉCNICOS')
run.bold = True
run.font.size = Pt(16)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('TT 2026-B066  —  Sistema FST con visión por computadora')
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(0x6B, 0x1F, 0x35)

doc.add_paragraph()

def add_entry(doc, term, que_es, facil, ejemplo):
    p = doc.add_paragraph()
    run = p.add_run(term)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6B, 0x1F, 0x35)
    p2 = doc.add_paragraph()
    lbl = p2.add_run('Qué es: ')
    lbl.bold = True
    lbl.font.size = Pt(10)
    p2.add_run(que_es).font.size = Pt(10)
    p3 = doc.add_paragraph()
    lbl2 = p3.add_run('Como para un niño de 5 años: ')
    lbl2.bold = True
    lbl2.font.size = Pt(10)
    r = p3.add_run(facil)
    r.font.size = Pt(10)
    r.font.italic = True
    p4 = doc.add_paragraph()
    lbl3 = p4.add_run('En el proyecto: ')
    lbl3.bold = True
    lbl3.font.size = Pt(10)
    p4.add_run(ejemplo).font.size = Pt(10)
    doc.add_paragraph()

def add_section(doc, title):
    p = doc.add_paragraph()
    run = p.add_run('── ' + title + ' ──')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    doc.add_paragraph()

# ── EL EXPERIMENTO ──
add_section(doc, 'EL EXPERIMENTO Y EL LABORATORIO')

add_entry(doc,
    'FST — Forced Swim Test (Prueba de Nado Forzado)',
    'Experimento preclínico donde se mete una rata en un cilindro con agua y se observa su comportamiento durante 5 minutos. Los investigadores miden cuánto tiempo pasa inmóvil, nadando o escalando para evaluar si un fármaco tiene efecto antidepresivo.',
    'Imagina que metes un ratoncito de juguete en una tina y cronometras cuánto tiempo nada, cuánto se queda quieto y cuánto intenta salirse. El tiempo que pasa quieto dice qué tan "triste" está.',
    'El laboratorio del Dr. Sandino hace este experimento para probar nuevas moléculas que podrían servir como antidepresivos antes de dárselas a personas.')

add_entry(doc,
    'Inmovilidad / Nado activo / Escalamiento',
    'Las tres conductas que el sistema detecta. Inmovilidad: la rata flota sin moverse. Nado activo: paladas horizontales. Escalamiento: movimientos rápidos de patas delanteras hacia las paredes del cilindro. Son mutuamente excluyentes: en cada instante la rata solo hace una.',
    'Imagina a alguien en una alberca: puede flotar sin moverse, nadar de un lado a otro, o intentar salirse por la orilla. En cada momento solo hace una de las tres cosas.',
    'El sistema clasifica cada frame del video con una etiqueta: inmovilidad, nado o escalamiento. Al final suma cuántos segundos pasó en cada estado por rata.')

add_entry(doc,
    'Desesperanza conductual (behavioral despair)',
    'Término científico que describe el momento en que la rata deja de luchar y se queda flotando. Se usa como indicador de un estado análogo a la depresión en el modelo preclínico.',
    'Es cuando la rata ya se cansó de intentar salir y simplemente flota. Los científicos dicen que eso se parece a cómo se siente una persona con depresión que ya no quiere intentarlo.',
    'Cuando el sistema mide mucho tiempo de inmovilidad, el investigador interpreta que el fármaco no tuvo efecto antidepresivo en ese espécimen.')

add_entry(doc,
    'Modelo preclínico',
    'Experimento realizado con animales de laboratorio (aquí ratas) antes de pasar a pruebas en humanos. Se usa para evaluar si un fármaco tiene el efecto esperado con menor riesgo ético.',
    'Antes de darle una medicina nueva a personas, los científicos la prueban con ratas para ver si funciona y si es peligrosa. Es como probar el fuego de una estufa nueva con una olla pequeña antes de cocinar el banquete.',
    'El FST es el modelo preclínico estándar para antidepresivos. Todos los antidepresivos que usamos hoy (como la fluoxetina) se probaron primero así.')

add_entry(doc,
    'Fluoxetina / Desipramina',
    'Fármacos antidepresivos usados como grupos de referencia en el experimento. La fluoxetina actúa sobre el sistema serotoninérgico (aumenta el nado activo); la desipramina sobre el noradrenérgico (aumenta el escalamiento). Ambos tienen efecto conocido y comprobado.',
    'Son las "respuestas correctas" del experimento: si el nuevo fármaco hace que la rata se comporte parecido a cuando tomó fluoxetina, es señal de que podría funcionar como antidepresivo.',
    'El grupo de referencia recibe fluoxetina para establecer cómo cambian las conductas con un antidepresivo conocido. El nuevo fármaco del grupo experimental se compara contra ese perfil.')

# ── VISIÓN POR COMPUTADORA ──
add_section(doc, 'VISIÓN POR COMPUTADORA')

add_entry(doc,
    'Visión por computadora (Computer Vision)',
    'Rama de la informática que enseña a las computadoras a interpretar imágenes y videos. La computadora analiza los píxeles para detectar objetos, movimientos o patrones sin intervención humana.',
    'Así como tus ojos ven y tu cerebro entiende lo que hay, la visión por computadora le enseña a la computadora a hacer lo mismo: "esto es una rata, esto es el cilindro".',
    'El sistema usa visión por computadora para ver el video del FST y detectar dónde está la rata y qué conducta hace, sin que nadie tenga que verlo a mano.')

add_entry(doc,
    'Frame (cuadro)',
    'Una fotografía individual dentro de un video. Un video de 5 minutos a 30 fps (frames por segundo) contiene 9,000 fotogramas. El sistema analiza cada uno por separado.',
    'Un video es como un flipbook: cada página es un frame. Cuando pasas las páginas rápido parece movimiento. Un video normal tiene 30 páginas por segundo.',
    'El pipeline clasifica la conducta de la rata en cada uno de los 9,000 frames del video de 5 minutos del Día 2.')

add_entry(doc,
    'CLAHE (Contrast Limited Adaptive Histogram Equalization)',
    'Técnica de procesamiento de imagen que mejora el contraste de forma adaptativa por zonas, sin sobreexponer las áreas que ya tienen buen contraste. Se activa cuando el video tiene reflejos o pelaje claro.',
    'Imagina una foto muy oscura donde usas el ajuste de brillo del teléfono. CLAHE hace lo mismo pero de forma inteligente: ajusta el brillo en cada zona por separado, sin quemar lo que ya se veía bien.',
    'Si el video del laboratorio tiene reflejos en el agua o pelaje muy claro, el sistema aplica CLAHE antes de analizar para que la detección sea más precisa.')

add_entry(doc,
    'Detección de objetos',
    'Proceso por el cual un algoritmo identifica y localiza objetos específicos dentro de una imagen, marcándolos con un recuadro (bounding box). Indica qué es el objeto y dónde está exactamente.',
    'Como buscar a Wally y ponerle un círculo cuando lo encuentras. La computadora hace lo mismo: busca la rata y dibuja un cuadrito alrededor indicando su posición exacta.',
    'YOLOv8 detecta los cilindros y las ratas en cada frame del video, junto con la confianza de detección.')

add_entry(doc,
    'Seguimiento multi-objeto (Multi-Object Tracking)',
    'Técnica que mantiene la identidad de cada objeto detectado a lo largo del tiempo en un video. Cuando hay varias ratas, el sistema sabe cuál es cuál en todo momento, aunque se muevan o se tapen.',
    'Imagina tres globos de colores en una fiesta: aunque se crucen y se muevan, siempre sabes cuál es el rojo, el azul y el verde. ByteTrack hace lo mismo con las ratas.',
    'ByteTrack garantiza que "Rata 1" siempre sea la misma rata en todos los frames del video, incluso después de una oclusión.')

add_entry(doc,
    'Confianza de detección',
    'Valor entre 0 y 1 que indica qué tan seguro está el modelo de que lo que detectó es realmente el objeto buscado. 0.95 = muy seguro; 0.30 = mucha duda.',
    'Como cuando alguien te pregunta si estás seguro de que eso es un gato y respondes "90% seguro". La computadora da un número que dice qué tan confiada está de su respuesta.',
    'El sistema solo acepta detecciones de cilindros con confianza mayor o igual a 0.70 (RN-11). Por debajo de eso genera un reporte de error de diagnóstico.')

add_entry(doc,
    'Oclusión',
    'Fenómeno donde un objeto tapa a otro en el video, haciéndolo temporalmente invisible. En el FST ocurre cuando una rata se sumerge o dos ratas se cruzan.',
    'Es como cuando alguien se para enfrente de ti y ya no puedes ver a la persona que estaba atrás. Para la computadora es como si la rata hubiera desaparecido por un momento.',
    'El filtro de Kalman predice dónde debería estar la rata durante la oclusión, y ByteTrack la reconoce cuando reaparece.')

# ── APRENDIZAJE SUPERVISADO ──
add_section(doc, 'APRENDIZAJE SUPERVISADO Y MODELOS')

add_entry(doc,
    'Aprendizaje supervisado (Supervised Learning)',
    'Tipo de aprendizaje automático donde se entrena un modelo mostrándole muchos ejemplos etiquetados. El modelo aprende los patrones y después clasifica ejemplos nuevos que nunca vio.',
    'Es como enseñarle a un niño a reconocer frutas: le muestras muchas fotos y le dices el nombre. Después el niño puede reconocer una fruta nueva él solo.',
    'El clasificador ResNet-18 se entrena con recortes de video anotados en BORIS. Después clasifica solos los frames nuevos del experimento.')

add_entry(doc,
    'Aprendizaje profundo (Deep Learning)',
    'Subconjunto del aprendizaje automático que usa redes neuronales con muchas capas. Puede aprender patrones muy complejos directamente de imágenes o video sin que un humano defina reglas manualmente.',
    'Imagina una cadena de amigos pasándose una pelota: cada amigo le da forma antes de pasarla al siguiente. Al final la pelota tiene la forma perfecta. Cada "amigo" en la cadena es una capa de la red neuronal.',
    'YOLOv8 y ResNet-18 son modelos de aprendizaje profundo con millones de conexiones que aprendieron a reconocer formas, texturas y movimientos.')

add_entry(doc,
    'Red neuronal',
    'Estructura matemática inspirada en el cerebro humano, formada por capas de "neuronas" artificiales conectadas. Cada conexión tiene un peso numérico que se ajusta durante el entrenamiento.',
    'Como una red de teléfonos conectados: la señal entra por un lado, pasa por muchos teléfonos que la transforman, y al final sale la respuesta. Cuantos más teléfonos (neuronas), más complejo lo que puede aprender.',
    'YOLOv8 y ResNet-18 son redes neuronales. ResNet-18 tiene 18 capas de transformaciones matemáticas entre la imagen de entrada y la etiqueta de conducta que produce.')

add_entry(doc,
    'YOLOv8 (You Only Look Once, versión 8)',
    'Modelo de detección de objetos muy rápido. Analiza toda la imagen de una sola pasada en lugar de escanearla varias veces, lo que lo hace eficiente para video en tiempo real.',
    'La mayoría de los detectores de objetos miran la imagen varias veces para estar seguros. YOLO es el detective que con una sola mirada ya sabe todo lo que hay en la imagen.',
    'El sistema usa YOLOv8 para localizar los cilindros y las ratas en cada frame. Si la confianza baja de 0.70, se activa el modo de diagnóstico.')

add_entry(doc,
    'ResNet-18 / ResNet-50 (Residual Network)',
    'Arquitecturas de redes neuronales para clasificar imágenes. El número indica la cantidad de capas. ResNet-18 es ligera y corre en CPU; ResNet-50 es más potente pero requiere más recursos.',
    'Dos cocineros: uno rápido para platillos sencillos (ResNet-18) y uno más lento para platillos más elaborados (ResNet-50). Primero pruebas con el rápido; si no resulta bien, llamas al otro.',
    'ResNet-18 clasifica cada frame en inmovilidad, nado o escalamiento. Si no alcanza F1 mayor o igual a 85 %, ResNet-50 actúa como segunda oportunidad.')

add_entry(doc,
    'Transfer Learning (Aprendizaje por transferencia)',
    'Técnica que toma un modelo ya entrenado para una tarea general y lo adapta para una tarea específica con mucho menos datos y tiempo de entrenamiento.',
    'Es como cuando ya sabes manejar coche automático y aprendes el manual: no empiezas de cero, ya sabes lo del volante y los espejos. Solo aprendes la parte nueva del clutch.',
    'ResNet-18 ya fue entrenado con ImageNet (1.2 millones de imágenes). Solo se ajusta la última capa para aprender a distinguir las tres conductas del FST.')

add_entry(doc,
    'ImageNet',
    'Base de datos masiva con 1.2 millones de imágenes etiquetadas en 1,000 categorías. Es el entrenamiento general estándar que usan casi todos los modelos de visión por computadora antes de especializarse.',
    'La escuela primaria para modelos de IA: ahí aprenden a reconocer miles de objetos. Cuando llegan al laboratorio ya vienen "educados" y solo necesitan aprender las tres conductas del FST.',
    'ResNet-18 usó ImageNet como entrenamiento inicial. Los conocimientos sobre bordes, texturas y formas son útiles para reconocer las posturas de la rata.')

add_entry(doc,
    'ByteTrack',
    'Algoritmo de seguimiento multi-objeto que asocia detecciones entre frames para mantener la identidad de cada objeto. Usa incluso detecciones de baja confianza, lo que reduce errores durante oclusiones.',
    'Como el juego de tres vasos y una bolita: aunque los muevan rápido y los crucen, ByteTrack siempre sabe en qué vaso está la bolita. Con ratas: aunque se sumerjan o se crucen, sabe cuál es cuál.',
    'ByteTrack mantiene la identidad de cada rata en todos los frames del video, incluso cuando se mueven rápido o quedan tapadas brevemente.')

add_entry(doc,
    'Filtro de Kalman',
    'Algoritmo matemático que predice dónde estará un objeto en el siguiente frame basándose en su posición y velocidad actuales. Ayuda al seguimiento cuando la detección falla momentáneamente.',
    'Es como adivinar dónde va a estar una pelota rodando: si la ves ir hacia la derecha, ya sabes que en un segundo estará más a la derecha aunque ya no la veas. El filtro hace esa predicción con matemáticas.',
    'Cuando una rata se sumerge y YOLOv8 no la detecta por 2-3 frames, el filtro de Kalman predice dónde debería estar y ByteTrack la recupera cuando reaparece.')

add_entry(doc,
    'F1, Precisión y Recall',
    'Métricas estadísticas para evaluar un clasificador. Precisión: de las veces que dijo "es inmovilidad", cuántas tenía razón. Recall: de todos los momentos de inmovilidad reales, cuántos detectó. F1: promedio balanceado entre ambos.',
    'Como un detector de mentiras: precisión sería "de las veces que dijo mentira, cuántas eran realmente mentira". Recall sería "de todas las mentiras que hubo, cuántas descubrió". F1 combina los dos en un solo número.',
    'El criterio de aceptación es F1 mayor o igual a 85 % por clase. Si escalamiento obtiene F1 del 70 %, el sistema aplica el fallback de RN-13.')

add_entry(doc,
    'Fallback',
    'Plan B que el sistema ejecuta automáticamente cuando el plan A no cumple los criterios de calidad. Si el clasificador no distingue bien escalamiento de nado, agrupa ambas conductas en "conducta activa".',
    'Como cuando se acaba el platillo del menú y el mesero te ofrece la alternativa. No es lo que pediste, pero todavía puedes comer. El fallback es la alternativa automática del sistema cuando el plan A falla.',
    'Si ResNet-18 y ResNet-50 no alcanzan F1 mayor o igual a 85 % distinguiendo escalamiento de nado, el reporte dice "conducta activa" (ambas juntas), equivalente al protocolo original de Porsolt 1977.')

add_entry(doc,
    'BORIS (Behavioral Observation Research Interactive Software)',
    'Software gratuito para anotar conductas en videos. El investigador ve el video y presiona teclas para marcar cuándo empieza y termina cada conducta, generando un archivo con los tiempos exactos.',
    'Como una app de cronómetro especial: ves el video y cuando la rata empieza a nadar aprietas N, cuando para la sueltas. Al final tienes un registro exacto de cuándo hizo qué.',
    'El equipo usa BORIS para etiquetar manualmente los videos de entrenamiento. Con esos datos etiquetados el modelo de clasificación aprende a reconocer cada conducta.')

add_entry(doc,
    'Dataset (conjunto de datos)',
    'Colección de ejemplos etiquetados para entrenar y evaluar un modelo. En el proyecto: recortes de video de ratas donde cada imagen tiene la etiqueta de qué conducta realizaban.',
    'Es el libro de ejercicios con respuestas: el modelo estudia ese libro para aprender. Después se le hace un examen con ejemplos nuevos que nunca vio.',
    'El dataset son frames de los videos del Dr. Sandino anotados en BORIS. El conjunto de prueba son videos separados, distintos a los de entrenamiento.')

add_entry(doc,
    'scikit-learn',
    'Biblioteca de Python para calcular métricas de clasificación. Aquí se usa para calcular F1, precisión y recall por clase después de que el modelo clasifica los frames.',
    'Una calculadora especializada que no suma números, sino que calcula qué tan bien le fue al modelo comparando sus respuestas con las respuestas correctas.',
    'Después de que ResNet-18 clasifica los frames del conjunto de prueba, scikit-learn calcula F1 por clase y determina si se cumple el criterio del 85 %.')

add_entry(doc,
    'PyTorch',
    'Biblioteca de Python para construir y entrenar modelos de aprendizaje profundo (redes neuronales). Es el framework nativo de YOLOv8 y el ecosistema de referencia para ResNet.',
    'Es como el taller donde se fabrican y reparan los modelos de IA: proporciona las herramientas para construir la red neuronal, entrenarla con datos y después usarla para clasificar.',
    'El sistema usa PyTorch como base para YOLOv8 (detección) y ResNet-18 (clasificación). Python 3.11 fue elegido por ser la versión con mejor compatibilidad con PyTorch.')

add_entry(doc,
    'OpenCV',
    'Biblioteca de código abierto para procesamiento de imágenes y video en Python. Decodifica los archivos de video, extrae los frames y aplica operaciones como CLAHE.',
    'Es la navaja suiza para trabajar con videos e imágenes en código: abre el archivo de video, lo parte en fotogramas individuales y los procesa uno por uno.',
    'OpenCV decodifica el video .mp4 que sube el investigador y entrega los frames a YOLOv8. También aplica CLAHE si el contraste del video es insuficiente.')

# ── LA PLATAFORMA WEB ──
add_section(doc, 'LA PLATAFORMA WEB')

add_entry(doc,
    'SPA (Single Page Application)',
    'Tipo de aplicación web que carga una sola página HTML y actualiza el contenido dinámicamente sin recargar. La navegación entre secciones es instantánea.',
    'Las páginas web viejas eran como libros: cada vez que ibas a otra página tenías que dar vuelta completa (recargar todo). Una SPA es una pizarra mágica: cambias el dibujo sin tener que borrar y redibujar todo.',
    'La interfaz del investigador es una SPA: navega entre subir video, ver resultados y el dashboard sin que la página se recargue cada vez.')

add_entry(doc,
    'API REST',
    'Forma estándar de comunicación entre el navegador del usuario (frontend) y el servidor (backend). El navegador hace peticiones HTTP y recibe respuestas en formato JSON.',
    'Tú eres el cliente en un restaurante (frontend), el mesero es la API, y la cocina es el servidor (backend). Le dices al mesero lo que quieres (petición), la cocina lo prepara y el mesero te lo trae (respuesta).',
    'Cuando el investigador hace clic en subir video, el frontend hace una petición al backend que responde "recibido, en cola".')

add_entry(doc,
    'JWT (JSON Web Token)',
    'Método de autenticación que usa un código cifrado para verificar quién es el usuario. Después del login el servidor entrega un token; el usuario lo adjunta en cada petición posterior.',
    'Como la pulsera que te ponen en la entrada de un parque de diversiones. Cada vez que quieres subir a un juego, enseñas la pulsera sin tener que comprar boleto de nuevo.',
    'Cuando el investigador hace login con su correo @ipn.mx, recibe un JWT. Ese token viaja en cada petición posterior para que el servidor verifique su identidad.')

add_entry(doc,
    'Backend',
    'Parte del sistema que corre en el servidor y que el usuario no ve. Gestiona la lógica, valida los datos, interactúa con la base de datos y coordina el pipeline.',
    'Es la cocina del restaurante: el comensal solo ve el plato, pero en la cocina hay cocineros preparando, cortando y sazonando. El backend es todo lo que pasa detrás de la pantalla.',
    'El backend de Flask recibe el video, verifica que sea .mp4, lo almacena y encola el análisis para el worker.')

add_entry(doc,
    'Worker asíncrono',
    'Proceso independiente que ejecuta tareas largas en segundo plano sin bloquear al servidor web. El worker toma el video de la cola y ejecuta el pipeline mientras el servidor atiende otras peticiones.',
    'En una cafetería, el barista puede atender nuevos clientes mientras el café se hace solo en la máquina. El worker es la máquina: trabaja solo para que el barista no tenga que quedarse parado esperando.',
    'El worker ejecuta el pipeline de análisis (puede tardar varios minutos) mientras el servidor sigue respondiendo otras peticiones de login o del dashboard.')

add_entry(doc,
    'Cola de tareas (encolado)',
    'Mecanismo que almacena tareas pendientes en orden para que los workers las procesen de una en una. Evita que el servidor colapse si llegan muchos trabajos simultáneos.',
    'La fila del banco con números: los clientes llegan, toman su número y esperan su turno. La cola de tareas es esa fila y el worker es el cajero que atiende de uno en uno.',
    'El sistema acepta hasta 4 usuarios simultáneos. Si los 4 suben un video al mismo tiempo, se forman en la cola y el worker los procesa secuencialmente.')

add_entry(doc,
    'PostgreSQL',
    'Sistema de base de datos relacional de código abierto. Almacena datos en tablas, garantiza integridad con propiedades ACID y usa SQL como lenguaje de consulta.',
    'Un archivero muy ordenado: cada experimento tiene su cajón, cada rata su carpeta, cada conducta su ficha. Puedes buscar y combinar información sin perder nada, aunque la aplicación se reinicie.',
    'PostgreSQL guarda usuarios, experimentos, resultados por conducta y reportes. Si el servidor se reinicia, todos los datos siguen ahí exactamente como estaban.')

add_entry(doc,
    'ACID (Atomicidad, Consistencia, Aislamiento, Durabilidad)',
    'Propiedades que garantizan que las operaciones en la base de datos son confiables. Si algo falla a la mitad, la base queda en estado consistente, sin datos a medias.',
    'Cuando transfieres dinero: ACID garantiza que si el internet se cae, o tu cuenta pierde el dinero Y se lo depositan al otro, o ninguna de las dos cosas ocurre. Nunca queda a medias.',
    'Si el worker falla mientras escribe los resultados de una rata, ACID garantiza que no quedan resultados parciales: o se guardaron todos o no se guardó ninguno.')

add_entry(doc,
    'Flask',
    'Microframework de Python para crear APIs y aplicaciones web. Es ligero y flexible, con extensiones para autenticación JWT y base de datos.',
    'Un kit de herramientas básico para construir el servidor de una web. Como un kit de LEGO: viene con las piezas esenciales y tú decides cómo armarlas para lo que necesitas.',
    'El backend del sistema está en Flask. Expone los endpoints de la API REST y coordina el worker asíncrono.')

add_entry(doc,
    'React.js + Vite',
    'React es la biblioteca de JavaScript para construir interfaces interactivas basadas en componentes reutilizables. Vite es la herramienta que compila el código React muy rápido durante el desarrollo.',
    'React son piezas de LEGO para construir páginas web (botones, formularios, gráficas). Vite es la fábrica que construye esas piezas rapidísimo para que puedas probarlas de inmediato.',
    'La interfaz del investigador (subir video, ver progreso, descargar resultados) está construida con React. Vite lo compila en segundos cada vez que el desarrollador hace un cambio.')

add_entry(doc,
    'React Router',
    'Biblioteca que gestiona la navegación entre diferentes vistas dentro de una SPA. Aunque la URL del navegador cambia, no hay recarga de página.',
    'El mapa de rutas de la aplicación. Como el GPS: aunque estés en la misma ciudad (misma página), te lleva a la dirección correcta (vista correcta) según a dónde quieres ir.',
    'React Router controla que cuando el investigador hace clic en Ver resultados, la URL cambie a /resultados y se muestre esa vista sin recargar la aplicación.')

add_entry(doc,
    'Docker / Docker Compose',
    'Docker empaqueta cada componente (frontend, backend, PostgreSQL, worker) en un contenedor aislado con todo lo que necesita. Docker Compose levanta todos los contenedores con un solo comando.',
    'Cada parte del sistema vive en su propia caja con su propia electricidad y herramientas. Docker Compose abre todas las cajas al mismo tiempo con un solo botón.',
    'Para instalar el sistema en el laboratorio solo se necesita Docker. Con el comando "docker compose up" se levantan los 4 contenedores sin instalar Python, Node.js ni nada más.')

add_entry(doc,
    'ESLint',
    'Herramienta de análisis estático de código JavaScript que detecta errores de estilo y posibles bugs antes de ejecutar el programa. Obliga al equipo a seguir las mismas reglas de escritura.',
    'El corrector de ortografía de Word, pero para código: te subraya en rojo cuando escribes algo incorrecto o cuando no sigues el estilo acordado. Así todo el código se ve igual sin importar quién lo escribió.',
    'Cuando Ángel escribe código JavaScript del frontend, ESLint le avisa si hay variables sin usar o si no siguió el estilo de indentación acordado.')

# ── INGENIERÍA DE SOFTWARE ──
add_section(doc, 'INGENIERÍA DE SOFTWARE')

add_entry(doc,
    'Scrum',
    'Marco de trabajo ágil que organiza el desarrollo en ciclos cortos llamados sprints (2 semanas). Permite adaptarse a cambios sin rehacer todo el diseño.',
    'En lugar de cocinar un banquete entero de una vez (arriesgándote a que todo salga mal), lo preparas plato por plato. Cada sprint terminas algo funcional y ajustas la receta del siguiente.',
    'En TT-I los sprints cubrieron análisis, diseño de arquitectura y documentación del protocolo. En TT-II cubrirán implementación y evaluación del sistema.')

add_entry(doc,
    'Sprint / Backlog',
    'Sprint: ciclo de desarrollo de 2 semanas con tareas definidas al inicio. Backlog: lista priorizada de todas las tareas pendientes del proyecto.',
    'El backlog es la lista del súper de todo lo que necesitas comprar. El sprint es la visita de hoy: decides qué vas a buscar en esta vuelta y lo que no conseguiste queda para la próxima.',
    'El backlog del proyecto incluye implementar login, integrar YOLOv8, diseñar el dashboard, etc. En cada sprint se toman las tareas de mayor prioridad.')

add_entry(doc,
    'Pipeline',
    'Secuencia de etapas donde la salida de cada una es la entrada de la siguiente: preprocesamiento, detección, seguimiento, clasificación, generación de resultados.',
    'Una línea de ensamblaje de coches: primero el chasis, luego el motor, luego las puertas, luego la pintura. Cada estación hace una cosa y pasa el resultado a la siguiente.',
    'El video entra al pipeline: CLAHE mejora el contraste, YOLOv8 detecta las ratas, ByteTrack las sigue, ResNet-18 clasifica la conducta y finalmente se generan los reportes.')

add_entry(doc,
    'RF (Requerimiento Funcional)',
    'Descripción de una función específica que el sistema debe hacer. Responde a la pregunta: qué debe hacer el sistema. Ejemplo: "El sistema debe aceptar archivos .mp4".',
    'Son las reglas del juego escritas en contrato: "el sistema DEBE poder hacer X". Como en una compraventa: "la casa DEBE tener dos baños". Si no los cumple, el sistema está incompleto.',
    'El sistema tiene 31 RF. RF-08 dice que solo acepta archivos .mp4. Si alguien intenta subir un .avi, el sistema lo rechaza desde el cliente (RF-10 / RN-10).')

add_entry(doc,
    'RNF (Requerimiento No Funcional)',
    'Describe qué tan bien debe comportarse el sistema: velocidad, disponibilidad, seguridad, usabilidad. No son funciones sino atributos de calidad.',
    'Si los RF son el "qué", los RNF son el "qué tan bien". Como en un restaurante: RF = traer la comida; RNF = que llegue caliente en menos de 15 minutos.',
    'RNF-02: F1 mayor o igual a 85 % por clase. RNF-03: disponibilidad mayor o igual a 95 % mensual. Están mapeados a atributos de ISO/IEC 25010.')

add_entry(doc,
    'Regla de negocio (RN)',
    'Restricción o condición que el sistema debe respetar para que el resultado tenga validez científica o administrativa. No son funciones sino políticas.',
    'Son las reglas del reglamento interno: no son lo que el sistema hace, sino lo que NO puede hacer o cómo debe comportarse en casos especiales.',
    'RN-07: solo una conducta activa por frame por rata (mutuamente excluyentes). RN-11: detección solo válida si confianza mayor o igual a 0.70.')

add_entry(doc,
    'ISO/IEC 25010',
    'Norma internacional que define los atributos de calidad del software: funcionalidad, rendimiento, usabilidad, confiabilidad, seguridad, mantenibilidad y portabilidad.',
    'El reglamento oficial de construcción de software: define qué significa que un sistema sea "bueno". Como el código de construcción de edificios, pero para programas.',
    'Los 13 RNF del sistema están mapeados a atributos de esta norma. El objetivo general dice que el sistema debe cumplirla.')

add_entry(doc,
    'ISO/IEC/IEEE 12207',
    'Norma internacional que define los procesos del ciclo de vida del software: desde la concepción hasta el retiro. Incluye cómo documentar y gestionar el desarrollo.',
    'El manual de instrucciones oficial para construir software de forma profesional: dice qué pasos seguir, qué documentar y cómo organizarse en un proyecto.',
    'La arquitectura fue diseñada conforme a ISO/IEC/IEEE 12207:2017 (OE-3). La separación en contenedores Docker satisface específicamente su atributo de portabilidad.')

add_entry(doc,
    'NOM-062-ZOO-1999',
    'Norma Oficial Mexicana que establece las especificaciones técnicas para la producción, cuidado y uso de los animales de laboratorio en México. El laboratorio del Dr. Sandino opera bajo esta norma.',
    'Es el reglamento oficial del gobierno mexicano que dice cómo deben tratarse los animales usados en experimentos científicos: qué espacio necesitan, cómo deben cuidarse, cuándo y cómo pueden usarse.',
    'El sistema no interactúa directamente con las ratas, pero se menciona la NOM-062 para dejar claro que los videos que analiza provienen de un laboratorio que opera bajo ese marco ético.')

add_entry(doc,
    'Diagrama de casos de uso',
    'Diagrama UML que muestra qué puede hacer cada actor (usuario, sistema externo) con el sistema. Los actores son figuras de palito y las acciones son elipses.',
    'Como la lista de superpoderes de cada personaje: "Batman puede lanzar batarangs, conducir el Batimóvil". Los casos de uso listan qué puede hacer cada tipo de usuario del sistema.',
    'Los actores del sistema son Investigador, Administrador y Sistema de análisis, con casos de uso organizados en 5 paquetes: autenticación, experimentos, análisis, resultados y administración.')

add_entry(doc,
    'Diagrama de secuencia',
    'Diagrama UML que muestra el orden de los mensajes entre componentes durante un proceso. Cada columna es un componente y las flechas muestran quién le habla a quién y cuándo.',
    'El guión de una obra de teatro: dice en qué orden habla cada personaje y qué le dice al otro. En software los "personajes" son el navegador, el servidor, la base de datos.',
    'La diapositiva 14 muestra: Frontend sube video al Backend, el Backend encola en el Worker, el Worker analiza y guarda en PostgreSQL, y notifica al Frontend al terminar.')

add_entry(doc,
    'Arquitectura modular',
    'Diseño de software donde cada función está en un componente independiente que puede desarrollarse, probarse y reemplazarse sin afectar a los demás.',
    'Un equipo de sonido: el tocadiscos, el amplificador y las bocinas son módulos separados. Si el amplificador se daña, lo cambias sin tirar el tocadiscos ni las bocinas.',
    'El pipeline tiene módulos separados: detección, seguimiento, clasificación. Si ByteTrack no funciona bien, se puede reemplazar por otro algoritmo sin tocar YOLOv8 ni ResNet-18.')

add_entry(doc,
    'Uptime / Disponibilidad',
    'Porcentaje del tiempo que un sistema está operativo y accesible. Un uptime de 95 % mensual significa que puede estar caído no más de 36 horas en todo el mes.',
    'Una tienda que dice estar abierta el 95 % del tiempo puede cerrar máximo 1.5 días por mes. Si cierra más, no cumple su promesa.',
    'RNF-03 establece disponibilidad mayor o igual a 95 % mensual como requerimiento de diseño, no como contrato de servicio permanente.')

add_entry(doc,
    'Reproducibilidad',
    'Propiedad científica que indica que un experimento produce los mismos resultados bajo las mismas condiciones. Fundamental para que los datos sean publicables.',
    'Si cada vez que haces la receta de tu abuela el pastel sabe diferente, la receta no es reproducible. En ciencia, si dos laboratorios hacen el mismo experimento y obtienen resultados distintos, hay un problema.',
    'El sistema tiene reproducibilidad perfecta: el mismo video siempre produce el mismo resultado. El proceso manual no: dos evaluadores registran tiempos distintos para la misma rata.')

add_entry(doc,
    'Conjunto de prueba independiente',
    'Subconjunto de datos que el modelo nunca vio durante el entrenamiento, usado solo para medir su desempeño real. Si se mezcla con el entrenamiento, los resultados serán artificialmente buenos.',
    'El examen final de la escuela: las preguntas no pueden ser las mismas del cuaderno de ejercicios. Si no, el alumno solo memorizó sin aprender. El conjunto de prueba son las preguntas nuevas.',
    'F1 mayor o igual a 85 % se medirá sobre videos que el Dr. Sandino no usó para entrenar. Así se garantiza que el modelo aprendió a generalizar.')

add_entry(doc,
    'Trazabilidad',
    'Capacidad de rastrear cada decisión, dato o resultado hasta su origen. Aquí: saber exactamente qué clasificó el modelo, con qué confianza y en qué frame.',
    'Como tener un historial médico completo: el doctor puede ver todas las consultas y qué medicamentos tomaste. Sin trazabilidad, no sabes de dónde viene un resultado.',
    'El sistema genera un reporte de diagnóstico con nivel de confianza por espécimen y un video anotado donde el Dr. Sandino puede revisar cuadro a cuadro qué clasificó el modelo.')

out = r'C:\Users\ID634\Desktop\proyectos\latexFST\fst-ai-system-latex\Glosario_TT_B066.docx'
doc.save(out)
print("OK:", out)
