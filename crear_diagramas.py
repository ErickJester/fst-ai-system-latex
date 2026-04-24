from docx import Document
from docx.shared import Mm, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import os

BASE_DIR = r"C:\Users\ID634\Desktop\proyectos\latexFST\fst-ai-system-latex\diagramas"
OUTPUT   = r"C:\Users\ID634\Desktop\proyectos\latexFST\fst-ai-system-latex\diagramas.docx"
MARGIN   = Cm(1)

# (archivo, titulo)
DIAGRAMS = [
    ("image1.jpg",  "Diagrama General del Sistema"),
    ("image2.png",  "Diagrama de Casos de Uso \u2014 Paquete 1: Autenticaci\u00f3n y Gesti\u00f3n de Usuarios"),
    ("image3.png",  "Diagrama de Casos de Uso \u2014 Paquete 2: Gesti\u00f3n de Experimentos y Carga de Video"),
    ("image4.png",  "Diagrama de Casos de Uso \u2014 Paquete 3: Pipeline de An\u00e1lisis Conductual"),
    ("image5.png",  "Diagrama de Casos de Uso \u2014 Paquete 4: Resultados y Reportes"),
    ("image6.png",  "Diagrama de Casos de Uso \u2014 Paquete 5: Dashboard y Administraci\u00f3n"),
    ("er.png",      "Diagrama Entidad-Relaci\u00f3n"),
    ("WhatsApp Image 2026-04-23 at 8.12.40 AM (1).jpeg", "Diagrama de Clases"),
    ("WhatsApp Image 2026-04-23 at 8.12.40 AM.jpeg",     "Diagrama de Estados del Video"),
    ("estado_sistema.png",                               "Diagrama de M\u00e1quina de Estados del Sistema"),
    ("WhatsApp Image 2026-04-23 at 8.12.40 AM (2).jpeg", "Diagrama de Actividades \u2014 Parte 1: Autenticaci\u00f3n, Crear Experimento y Carga de Video"),
    ("WhatsApp Image 2026-04-23 at 8.12.40 AM (3).jpeg", "Diagrama de Actividades \u2014 Parte 2: Pipeline de IA, Resultados y Reportes"),
    ("seq_login.png",            "Diagrama de Secuencia: Login (Investigador)"),
    ("seq_login_admin.png",      "Diagrama de Secuencia: Login (Administrador)"),
    ("seq_logout.png",           "Diagrama de Secuencia: Logout"),
    ("seq_registro.png",         "Diagrama de Secuencia: Registro de Usuario"),
    ("seq_cambio_pass_inv.png",  "Diagrama de Secuencia: Cambio de Contrase\u00f1a (Investigador)"),
    ("seq_cambio_pass_admin.png","Diagrama de Secuencia: Cambio de Contrase\u00f1a (Administrador)"),
    ("seq_gestion_usuarios.png", "Diagrama de Secuencia: Gesti\u00f3n de Usuarios"),
    ("seq_perfil.png",           "Diagrama de Secuencia: Actualizaci\u00f3n de Perfil"),
    ("seq_analisis.png",         "Diagrama de Secuencia: An\u00e1lisis de Video"),
    ("seq_progreso.png",         "Diagrama de Secuencia: Progreso de An\u00e1lisis"),
    ("seq_error.png",            "Diagrama de Secuencia: Error en Pipeline"),
    ("seq_notificaciones.png",   "Diagrama de Secuencia: Notificaciones"),
    ("seq_resultados.png",       "Diagrama de Secuencia: Consulta de Resultados"),
    ("seq_reportes.png",         "Diagrama de Secuencia: Generaci\u00f3n de Reportes"),
]

TITLE_HEIGHT = Cm(1.5)  # space reserved for the title


def configure_section(section, landscape: bool):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width  = Mm(297)
        section.page_height = Mm(210)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width  = Mm(210)
        section.page_height = Mm(297)
    section.left_margin   = MARGIN
    section.right_margin  = MARGIN
    section.top_margin    = MARGIN
    section.bottom_margin = MARGIN


def content_area(landscape: bool):
    if landscape:
        w = Mm(297) - 2 * MARGIN
        h = Mm(210) - 2 * MARGIN - TITLE_HEIGHT
    else:
        w = Mm(210) - 2 * MARGIN
        h = Mm(297) - 2 * MARGIN - TITLE_HEIGHT
    return w, h


def fit_image(img_w_px, img_h_px, max_w_emu, max_h_emu):
    ratio = img_w_px / img_h_px
    max_ratio = max_w_emu / max_h_emu
    if ratio >= max_ratio:
        return max_w_emu, int(max_w_emu / ratio)
    else:
        return int(max_h_emu * ratio), max_h_emu


def add_title(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    # remove space before/after
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(6)
    return para


doc = Document()

# Remove default empty paragraph
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)

first = True
for filename, title in DIAGRAMS:
    img_path = os.path.join(BASE_DIR, filename)
    if not os.path.exists(img_path):
        print(f"  SKIP (not found): {filename}")
        continue

    with Image.open(img_path) as img:
        img_w, img_h = img.size

    landscape = img_w > img_h

    if first:
        section = doc.sections[0]
        first = False
    else:
        section = doc.add_section()

    configure_section(section, landscape)

    cw, ch = content_area(landscape)
    fw, fh = fit_image(img_w, img_h, cw, ch)

    add_title(doc, title)

    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(0)
    img_para.paragraph_format.space_after  = Pt(0)
    img_para.add_run().add_picture(img_path, width=fw, height=fh)

    print(f"  OK: {title} ({'landscape' if landscape else 'portrait'})")

doc.save(OUTPUT)
print(f"\nGuardado en: {OUTPUT}")
