# gen_guion.py — Guión de defensa oral TT 2026-B066
# Genera Guion_Defensa_TT_2026-B066.docx con python-docx

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import os

OUT = os.path.join(os.path.dirname(__file__), "Guion_Defensa_TT_2026-B066.docx")

# ── Colores ──────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x1B, 0x3E)
TEAL   = RGBColor(0x00, 0xB4, 0xA0)
GREEN  = RGBColor(0x00, 0x6B, 0x5D)
GOLD   = RGBColor(0xB8, 0x86, 0x0B)
GRAY   = RGBColor(0x22, 0x22, 0x22)
LGRAY  = RGBColor(0x88, 0x88, 0x88)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ── python-docx helpers ──────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def set_para_border(para, side='bottom', color='00B4A0', size=6, space=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bdr = OxmlElement(f'w:{side}')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), str(size))
    bdr.set(qn('w:space'), str(space))
    bdr.set(qn('w:color'), color)
    pBdr.append(bdr)
    pPr.append(pBdr)

def set_table_width(table, width_dxa):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def set_col_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def run_add(para, text, bold=False, italic=False, color=None, size=11):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    r.font.size = Pt(size)
    r.font.name = 'Arial'
    return r

def para_spacing(para, before=0, after=80, line_rule=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before / 20)
    pf.space_after  = Pt(after  / 20)
    if line_rule:
        pf.line_spacing_rule = line_rule

# ── Funciones de contenido ────────────────────────────────────
def add_slide_heading(doc, num, title):
    p = doc.add_paragraph()
    para_spacing(p, before=20, after=4)
    set_para_border(p, 'bottom', '00B4A0', size=6, space=4)
    run_add(p, f"Diapositiva {num} — ", bold=True, color=LGRAY, size=13)
    run_add(p, title,                   bold=True, color=NAVY,  size=13)
    return p

def add_tiempo(doc, seg):
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=4)
    run_add(p, f"⏱ ~{seg} segundos", italic=True, color=LGRAY, size=9)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=4)
    run_add(p, f"[{text}]", italic=True, color=LGRAY, size=9)
    return p

def add_dialogo(doc, speaker_label, speaker_color, text, before=4):
    p = doc.add_paragraph()
    para_spacing(p, before=before, after=4)
    run_add(p, speaker_label + "  ", bold=True, color=speaker_color, size=11)
    run_add(p, text, color=GRAY, size=11)
    return p

def A(doc, text, before=4):
    return add_dialogo(doc, "(Á) Ángel:", NAVY, text, before)

def V(doc, text, before=4):
    return add_dialogo(doc, "(V) Vanesa:", GREEN, text, before)

def AV(doc, text, before=4):
    return add_dialogo(doc, "(Á+V) Ambos:", GOLD, text, before)

# ── Crear documento ───────────────────────────────────────────
doc = Document()

# Márgenes 2.5 cm
section = doc.sections[0]
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Estilos base
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# ── ENCABEZADO ────────────────────────────────────────────────
from docx.oxml.shared import OxmlElement as OE
hdr = section.header
hp = hdr.paragraphs[0]
hp.clear()
run_add(hp, "Guión de Defensa · TT 2026-B066 · ESCOM-IPN", color=LGRAY, size=9)
set_para_border(hp, 'bottom', '00B4A0', size=4, space=4)

# ── PIE DE PÁGINA con número ─────────────────────────────────
ftr = section.footer
fp = ftr.paragraphs[0]
fp.clear()
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_para_border(fp, 'top', '00B4A0', size=4, space=4)
run_add(fp, "Pág. ", color=LGRAY, size=9)
# Número de página
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run = fp.add_run()
run.font.size = Pt(9)
run.font.color.rgb = LGRAY
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

# ══════════════════════════════════════════════════════════════
# PORTADA DEL DOCUMENTO
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(p, before=40, after=8)
run_add(p, "GUIÓN DE DEFENSA ORAL", bold=True, color=NAVY, size=22)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(p, before=0, after=4)
run_add(p, "Trabajo Terminal 2026-B066", bold=True, color=TEAL, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(p, before=0, after=20)
run_add(p, "Prototipo de un sistema web de análisis conductual asistido por IA\nen un modelo de conducta depresiva en rata mediante nado forzado",
        italic=True, color=GRAY, size=11)

# Tabla de info
# Ancho total aprox. 16 cm  → 9072 dxa
TW = 9072
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(table, TW)

COL1, COL2 = 2520, TW - 2520
data = [
    ("Expositores",
     ["(Á) Ángel Ali Frausto Robles", "(V) Vanesa Rodríguez Verdín Sandoval"],
     ["0D1B3E", "006B5D"]),
    ("Institución",   ["ESCOM-IPN · Mayo 2026"], [None]),
    ("Directores",    ["Dr. Israel Salas Ramírez", "Dra. Martha Rosa Cordero López"], [None, None]),
    ("Tiempo total",  ["≈ 18–20 minutos (22 diapositivas)"], [None]),
]
bgs = ["EBF4FF", "F5F5F5", "EBF4FF", "F5F5F5"]

for i, (label, lines, colors) in enumerate(data):
    row = table.rows[i]
    c0, c1 = row.cells[0], row.cells[1]
    set_col_width(c0, COL1); set_col_width(c1, COL2)
    set_cell_bg(c0, bgs[i]); set_cell_bg(c1, bgs[i])
    set_cell_margins(c0, 80, 80, 160, 100)
    set_cell_margins(c1, 80, 80, 120, 120)
    c0.paragraphs[0].clear()
    run_add(c0.paragraphs[0], label, bold=True, color=NAVY, size=11)
    for j, line in enumerate(lines):
        p2 = c1.paragraphs[0] if j == 0 else c1.add_paragraph()
        p2.clear()
        col = RGBColor(*bytes.fromhex(colors[j])) if colors[j] else GRAY
        run_add(p2, line, bold=(colors[j] is not None), color=col, size=11)

doc.add_paragraph()  # Espacio antes del salto

# Salto de página
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# INSTRUCCIONES DE USO
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
para_spacing(p, before=8, after=6)
set_para_border(p, 'bottom', '0D1B3E', size=6, space=4)
run_add(p, "Instrucciones de uso", bold=True, color=NAVY, size=16)

for label, color, desc in [
    ("(Á) Ángel",     NAVY,  "  habla los fragmentos marcados con su etiqueta."),
    ("(V) Vanesa",    GREEN, "  habla los fragmentos marcados con su etiqueta."),
    ("(Á+V) Ambos",   GOLD,  "  hablan simultáneamente (saludos y cierre)."),
]:
    p2 = doc.add_paragraph()
    para_spacing(p2, before=2, after=2)
    run_add(p2, label, bold=True, color=color, size=11)
    run_add(p2, desc, color=GRAY, size=11)

p2 = doc.add_paragraph()
para_spacing(p2, before=2, after=2)
run_add(p2, "[Nota de escena]", italic=True, color=LGRAY, size=10)
run_add(p2, "  Indica una acción o transición física.", color=GRAY, size=11)

p2 = doc.add_paragraph()
para_spacing(p2, before=2, after=10)
run_add(p2, "⏱ ~XX segundos", italic=True, color=LGRAY, size=10)
run_add(p2, "  Tiempo estimado de habla para esa diapositiva.", color=GRAY, size=11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# GUIÓN — 22 DIAPOSITIVAS
# ══════════════════════════════════════════════════════════════

# ── DIAPOSITIVA 1 — Portada ───────────────────────────────────
add_slide_heading(doc, 1, "Portada")
add_tiempo(doc, 45)
add_note(doc, "Ambos de pie, uno a cada lado del podio. Ángel habla primero.")
A(doc, "Buenos días, sinodales. Mi nombre es Ángel Ali Frausto Robles.")
V(doc, "Y yo soy Vanesa Rodríguez Verdín Sandoval. En nombre de ambos, agradecemos la oportunidad de presentar nuestro Trabajo Terminal número 2026-B066: «Prototipo de un sistema web de análisis conductual asistido por inteligencia artificial en un modelo de conducta depresiva en rata mediante nado forzado», desarrollado en la Escuela Superior de Cómputo del IPN, bajo la dirección del Doctor Israel Salas Ramírez y la Doctora Martha Rosa Cordero López.")

# ── DIAPOSITIVA 2 — Agenda ────────────────────────────────────
add_slide_heading(doc, 2, "Agenda")
add_tiempo(doc, 40)
A(doc, "La presentación está organizada en seis secciones. Iniciamos con el contexto y la motivación del problema; pasamos al marco del TT —objetivos y justificación—; luego al fundamento teórico; enseguida a la ingeniería de requisitos; después al diseño del sistema; y cerramos con la validación, las conclusiones y el trabajo futuro.")

# ── DIAPOSITIVA 3 — Introducción ─────────────────────────────
add_slide_heading(doc, 3, "Introducción")
add_tiempo(doc, 75)
add_note(doc, "Vanesa toma la palabra.")
V(doc, "La depresión afecta a más de 280 millones de personas en el mundo. Para evaluar el potencial antidepresivo de nuevas moléculas antes de pasar a ensayos humanos, la comunidad científica recurre al modelo del nado forzado —el Forced Swim Test o FST—, propuesto por Porsolt y colaboradores en 1977. En este paradigma se colocan ratas individualmente en un cilindro con agua del que no pueden salir; al principio intentan escapar, pero tras el estrés acumulado adoptan una postura de inmovilidad pasiva, lo que se conoce como desesperanza conductual o behavioral despair. Los fármacos antidepresivos reducen ese tiempo de inmovilidad. Desde el año 2000 se han publicado más de ocho mil artículos con este paradigma, lo que confirma su validez predictiva y su vigencia como estándar en neurociencia preclínica.")

# ── DIAPOSITIVA 4 — Planteamiento del problema ───────────────
add_slide_heading(doc, 4, "Planteamiento del problema")
add_tiempo(doc, 70)
V(doc, "El Laboratorio de Bioquímica Estructural, Sección de Posgrado, de la ENMyH-IPN, dirigido por el Doctor Sandino Reyes López, registra estos experimentos grabando cuatro ratas simultáneamente en vista lateral. Cada sesión dura entre 5 y 20 minutos, y hoy el análisis se realiza de manera manual: cronómetro en mano, cuadro a cuadro. Esto genera cuatro problemas concretos.")
A(doc, "Primero, el tiempo: analizar un experimento completo toma varias horas por rata y sesión. Segundo, la variabilidad: dos evaluadores pueden usar criterios distintos al clasificar una misma conducta, comprometiendo la reproducibilidad. Tercero, no existe registro centralizado; los resultados quedan dispersos en hojas de cálculo y notas físicas. Y cuarto, la granularidad: un desglose minuto a minuto —relevante para la dinámica farmacológica— es inviable a mano.")

# ── DIAPOSITIVA 5 — Propuesta de solución ────────────────────
add_slide_heading(doc, 5, "Propuesta de solución")
add_tiempo(doc, 80)
A(doc, "Nuestra propuesta integra dos componentes. El primero es un pipeline de análisis de video que detecta automáticamente a cada rata dentro de su cilindro y clasifica cuadro a cuadro su conducta. Utiliza sustracción de fondo con CLAHE para corregir contraste, YOLOv8 para la detección, ByteTrack con filtro de Kalman para el seguimiento multi-objeto y ResNet-18 con transfer learning para la clasificación. Si el equipo no cuenta con GPU, el sistema dispone de un respaldo clásico con CSRT o KCF.")
V(doc, "El segundo componente es la plataforma web que permite a un investigador subir sus videos desde el navegador y recibir reportes cuantitativos de manera automática. Incluye carga de videos para el Día 1 y el Día 2, un dashboard con progreso en tiempo real, resultados desglosados por rata y por minuto, reportes descargables en PDF y CSV, y gestión de roles para administradores e investigadores.")

# ── DIAPOSITIVA 6 — Objetivos ─────────────────────────────────
add_slide_heading(doc, 6, "Objetivos")
add_tiempo(doc, 75)
add_note(doc, "Ángel retoma la exposición.")
A(doc, "El objetivo general es diseñar, desarrollar y validar un prototipo de sistema web asistido por IA que automatice la detección y cuantificación de las conductas del FST en ratas, con resultados comparables a la evaluación manual y cumpliendo los atributos de calidad de ISO/IEC 25010. Para lograrlo planteamos seis objetivos específicos: revisar el estado del arte en análisis automatizado del FST; documentar el protocolo experimental y las condiciones de grabación; diseñar la arquitectura conforme a ISO/IEC/IEEE 12207; implementar el pipeline de visión y el clasificador supervisado; validar con la kappa de Cohen y el MAE respecto al gold standard; y cuantificar la reducción de tiempo frente al análisis manual.")

# ── DIAPOSITIVA 7 — Justificación y Alcance ──────────────────
add_slide_heading(doc, 7, "Justificación y Alcance")
add_tiempo(doc, 75)
V(doc, "El sistema está justificado por cuatro razones: reduce las horas de trabajo manual y homogeniza criterios entre evaluadores; es completamente abierto y desplegable con Docker conforme a ISO/IEC/IEEE 12207; su clasificador es modular y podría adaptarse a otros paradigmas conductuales; y sus atributos de calidad son verificables con métricas concretas bajo ISO/IEC 25010. En cuanto al alcance, el sistema analiza videos del FST del laboratorio en vista lateral con cuatro cilindros y clasifica las tres conductas mutuamente excluyentes generando reportes comparativos Día 1 versus Día 2. Quedan fuera del alcance el análisis en tiempo real, la corrección automática de iluminación, la intervención con ratas —bajo NOM-062— y el despliegue institucional definitivo.")

# ── DIAPOSITIVA 8 — Estado del Arte ──────────────────────────
add_slide_heading(doc, 8, "Estado del Arte")
add_tiempo(doc, 70)
A(doc, "Al revisar las herramientas existentes encontramos que las soluciones comerciales como EthoVision XT y ANY-maze son propietarias y no están especializadas en FST. Las herramientas abiertas como ezTrack o DBscorer están parcialmente especializadas, pero no clasifican las tres conductas ni usan aprendizaje profundo. Los trabajos académicos recientes —DSAAN en 2024 y 3D-RCNN en 2025— aplican deep learning pero no se distribuyen como sistemas desplegables. Nuestra propuesta es la única que combina: especialización para FST, las tres conductas, deep learning, interfaz web, software abierto y rastreo multi-objeto con ByteTrack.")

# ── DIAPOSITIVA 9 — Marco Teórico Conductual ─────────────────
add_slide_heading(doc, 9, "Marco Teórico Conductual")
add_tiempo(doc, 65)
V(doc, "El FST define tres conductas mutuamente excluyentes, cada una asociada a un sistema neurotransmisor distinto. La inmovilidad —el espécimen flota con el mínimo movimiento para no hundirse— es el indicador principal de desesperanza conductual. El nado activo —paladas horizontales que desplazan a la rata— está asociado al sistema serotoninérgico y es el que reducen los ISRS como la fluoxetina. Los intentos de escape —movimientos rápidos hacia las paredes del cilindro— se asocian al sistema noradrenérgico, como en el caso de la desipramina. Distinguir estas tres conductas permite inferir sobre qué mecanismo molecular actúa el fármaco evaluado.")

# ── DIAPOSITIVA 10 — Marco Teórico Visión/DL ─────────────────
add_slide_heading(doc, 10, "Marco Teórico de Visión y Aprendizaje Profundo")
add_tiempo(doc, 90)
A(doc, "El pipeline se apoya en cuatro conceptos clave. La sustracción de fondo por mediana pixel a pixel modela el fondo estático de los primeros cuadros; CLAHE —Contrast Limited Adaptive Histogram Equalization— corrige el contraste local antes del umbralado de Otsu, y la morfología por apertura elimina el ruido residual.")
A(doc, "ByteTrack usa un filtro de Kalman para predecir la posición de la rata y aplica asociación en dos pasos: primero detecciones de alta confianza, luego las de baja confianza, para recuperar trayectorias incluso en casos de oclusión parcial.")
A(doc, "ResNet-18 es una red residual cuyos bloques de salto aprenden F(x) = H(x) − x, lo que permite que el gradiente fluya sin atenuación durante el entrenamiento de redes profundas.")
A(doc, "Y el transfer learning nos permite partir de los pesos preentrenados en ImageNet —con aproximadamente un millón de imágenes y mil clases—, congelar las capas de extracción de características y reentrenar únicamente la capa final con tres salidas, una por conducta, usando cuadros etiquetados manualmente en BORIS.")

# ── DIAPOSITIVA 11 — Metodología ─────────────────────────────
add_slide_heading(doc, 11, "Metodología")
add_tiempo(doc, 65)
V(doc, "Para el desarrollo adoptamos Scrum, una metodología ágil que organiza el trabajo en sprints de dos semanas. Cada sprint inicia con la planeación del incremento entregable; se realizan dailies para sincronización y desbloqueo temprano; al final hay una revisión con el director y con el laboratorio; y la retrospectiva permite ajustar criterios para el siguiente ciclo. En cuanto a la distribución: Ángel se concentra en el pipeline de visión —detección con YOLOv8, seguimiento con ByteTrack, clasificación con ResNet-18 y las pruebas de validación—.")
A(doc, "Y Vanesa tiene a cargo la plataforma web: el backend en Flask con JWT, el modelado de la base de datos, el worker asíncrono, el dashboard y los módulos de reportes. El análisis de requisitos, el diseño de la arquitectura y la integración completa del sistema los desarrollamos de manera conjunta.")

# ── DIAPOSITIVA 12 — Requerimientos Funcionales ───────────────
add_slide_heading(doc, 12, "Requerimientos Funcionales")
add_tiempo(doc, 90)
add_note(doc, "Ángel recorre los seis módulos señalando en pantalla.")
A(doc, "Definimos 29 requerimientos funcionales en seis módulos. Autenticación —RF-01 a RF-04—: login con correo institucional del IPN, logout con JWT, recuperación de contraseña y gestión de roles. Gestión de usuarios —RF-05 a RF-07—: el administrador crea, modifica y desactiva cuentas; el investigador actualiza su propio perfil.")
A(doc, "Carga de video —RF-08 a RF-12—: acepta únicamente archivos MP4, valida en cliente y servidor, registra metadatos y encola automáticamente el análisis. Análisis conductual —RF-13 a RF-17—: ejecuta el pipeline en cuatro etapas con progreso en tiempo real, detecta cilindros con confianza mínima de 0.70, clasifica cuadro a cuadro de forma mutuamente excluyente y genera un PDF de diagnóstico ante errores.")
A(doc, "Resultados —RF-18 a RF-21—: tiempo y porcentaje por conducta y rata, desglose por minuto, comparación entre días y exportación en PDF y CSV. Dashboard de administración —RF-22 a RF-29—: gestión de experimentos, avisos de expiración de videos a 30 días y monitoreo de disco, cola y análisis activos.")

# ── DIAPOSITIVA 13 — RNF y Reglas de Negocio ─────────────────
add_slide_heading(doc, 13, "Requerimientos No Funcionales y Reglas de Negocio")
add_tiempo(doc, 80)
V(doc, "Los 13 requerimientos no funcionales cubren los atributos de calidad de ISO/IEC 25010. Rendimiento: navegación menor a 3 segundos en el percentil 95. Disponibilidad del 95 % en horario laboral. Seguridad con JWT, bcrypt con salt y aislamiento por usuario. Usabilidad: un investigador sin experiencia técnica debe completar el flujo sin asistencia. Compatibilidad con Chrome, Firefox y Edge. Portabilidad con despliegue único en Docker Compose. Y mantenibilidad bajo PEP-8, ESLint y documentación por módulo.")
V(doc, "Las 13 reglas de negocio formalizan la lógica del protocolo. Las más relevantes: un solo video por experimento y máximo dos —Día 1 y Día 2—; los videos se borran a los 30 días con aviso de 7 días de anticipación, pero los resultados se conservan indefinidamente conforme a RN-06; y si la kappa de Cohen no supera 0.80, el sistema aplica el fallback combinando nado activo y escalamiento en la categoría «conducta activa».")

# ── DIAPOSITIVA 14 — Riesgos y Factibilidad ──────────────────
add_slide_heading(doc, 14, "Análisis de Riesgos y Factibilidad")
add_tiempo(doc, 70)
A(doc, "Identificamos cinco riesgos. R-01 y R-02 son de nivel alto: videos con reflejos o pelaje claro que bajen la confianza de detección —mitigado con CLAHE y reportes de diagnóstico—; y dataset insuficiente para alcanzar la kappa objetivo —mitigado con anotación en BORIS, data augmentation y el fallback de conducta activa.")
A(doc, "Los tres riesgos de nivel medio son: que el laboratorio no entregue los registros a tiempo —calendario pactado desde el inicio de TT2—; que el tiempo de análisis sea inaceptablemente largo —medición por etapa y evaluación de GPU—; y que la integración del worker asíncrono sea compleja —contratos de API y pruebas end-to-end.")
A(doc, "En cuanto a la factibilidad: técnicamente todo el stack es software libre y maduro —OpenCV, YOLOv8, PyTorch, Flask, PostgreSQL—; económicamente el costo de licencias es cero; y operativamente el despliegue con Docker Compose usa el equipo ya existente del laboratorio.")

# ── DIAPOSITIVA 15 — Arquitectura ────────────────────────────
add_slide_heading(doc, 15, "Diseño de Arquitectura")
add_tiempo(doc, 65)
add_note(doc, "Ángel señala las capas en el diagrama proyectado.")
A(doc, "La arquitectura sigue el patrón cliente-servidor en tres capas más un cuarto contenedor asíncrono. React con Vite gestiona la interfaz en el navegador. Flask con JWT es la API REST que valida las peticiones y las encola. PostgreSQL proporciona la persistencia con garantías ACID. Y el worker ejecuta el pipeline de IA de forma asíncrona, independiente del servidor web. El sistema está organizado en cinco paquetes: autenticación, carga de video, pipeline de IA, resultados y reportes, y dashboard de administración. Los actores son el investigador como actor principal, el administrador y el worker como actor de sistema.")

# ── DIAPOSITIVA 16 — Pipeline de IA ──────────────────────────
add_slide_heading(doc, 16, "Diseño del Pipeline de IA")
add_tiempo(doc, 80)
A(doc, "El pipeline ejecuta cuatro etapas secuenciales. La primera es el preprocesamiento: se decodifica el video, se aplica CLAHE si el contraste cae bajo el umbral configurable y se limpia el ruido con morfología de apertura.")
A(doc, "La segunda es la detección de cilindros: YOLOv8 localiza los cuatro ROIs con confianza mínima de 0.70; si falla, se genera un reporte de diagnóstico en PDF con la imagen del cuadro problemático.")
A(doc, "La tercera es el seguimiento: YOLOv8 detecta a las ratas cuadro a cuadro y ByteTrack con filtro de Kalman mantiene la identidad de cada rata a lo largo del video; el respaldo CSRT o KCF opera si no hay GPU disponible.")
A(doc, "La cuarta es la clasificación: ResNet-18 con transfer learning desde ImageNet asigna a cada cuadro una de las tres conductas. ResNet-50 está disponible como alternativa de mayor precisión si el tiempo de inferencia lo permite.")

# ── DIAPOSITIVA 17 — Base de Datos ───────────────────────────
add_slide_heading(doc, 17, "Diseño de la Base de Datos")
add_tiempo(doc, 65)
V(doc, "El modelo de datos está implementado en PostgreSQL con SQLAlchemy como ORM. Las entidades principales son: USERS para autenticación y roles; EXPERIMENTS como contenedor del experimento; VIDEOS con la referencia al archivo y la fecha de eliminación programada; JOBS que liga el video con el pipeline y conserva el progreso por etapa; ANIMALS con los ROIs por sujeto —hasta cuatro por video—; BEHAVIOR_RESULTS con los resultados de clasificación —nado, inmovilidad y escalamiento—; REPORTS para los archivos en PDF y CSV; y NOTIFICATIONS para los avisos de análisis completado, expiración y errores.")
V(doc, "El diseño cumple las reglas de negocio clave: eliminar un experimento elimina en cascada todo lo asociado, pero los resultados en BEHAVIOR_RESULTS se conservan aunque se borre el video, conforme a RN-06.")

# ── DIAPOSITIVA 18 — Vistas del Prototipo ────────────────────
add_slide_heading(doc, 18, "Vistas del Prototipo")
add_tiempo(doc, 75)
V(doc, "El prototipo contempla seis vistas principales. El Login con validación de correo institucional y mensaje de error genérico según RN-01. El Dashboard del investigador con estado de todos sus experimentos y avisos de expiración. La Creación de experimento con un formulario de tres pasos: metadatos, asignación de ROIs y carga de videos.")
V(doc, "La pantalla de Progreso del análisis con barra de avance por etapa del pipeline y manejo de error con reporte. La pantalla de Resultados con cuatro pestañas: resumen, detalle por rata, desglose por minuto y comparación Día 1 versus Día 2. Y el Panel de administración con monitoreo de disco, cola de trabajos y gestión de usuarios, con alertas al 80 % de capacidad.")

# ── DIAPOSITIVA 19 — Validación y Métricas ───────────────────
add_slide_heading(doc, 19, "Validación y Métricas")
add_tiempo(doc, 80)
add_note(doc, "Ángel retoma la exposición.")
A(doc, "La validación se realiza contra un gold standard construido con BORIS —Behavioral Observation Research Interactive Software—, donde expertos del laboratorio anotan manualmente los videos. Utilizamos tres métricas principales: la kappa de Cohen, que debe superar 0.80 para considerar la concordancia inter-evaluador como satisfactoria —en caso contrario se activa el fallback de conducta activa—; el F1-score por clase, que balancea precisión y recall para detectar sesgos en conductas minoritarias; y el error absoluto medio en segundos, que cuantifica la diferencia entre el tiempo calculado por el sistema y el registrado manualmente.")
A(doc, "Adicionalmente evaluaremos usabilidad y comportamiento temporal: al menos un usuario real del laboratorio debe completar el flujo sin asistencia técnica con una tasa de éxito del 90 %, y el tiempo de procesamiento debe ser menor o igual al tiempo de duración del video en el equipo del laboratorio, incluso sin GPU.")

# ── DIAPOSITIVA 20 — Conclusiones y Trabajo Futuro ───────────
add_slide_heading(doc, 20, "Conclusiones y Trabajo Futuro")
add_tiempo(doc, 80)
V(doc, "Como conclusiones del TT1 podemos afirmar: documentamos el protocolo del laboratorio y cuantificamos el costo real del análisis manual; el estado del arte confirmó la brecha —no existe ningún sistema abierto, web, especializado en FST que distinga las tres conductas—; definimos 29 requerimientos funcionales, 13 no funcionales y 13 reglas de negocio alineados a ISO/IEC 25010; y la arquitectura y el pipeline de IA propuestos son técnica y económicamente factibles con costo cero en licencias.")
A(doc, "Para el TT2, los pasos siguientes son: implementar el backend Flask con JWT, la base de datos y la API REST; desarrollar el frontend en React con Vite, el dashboard y los módulos de reportes; ampliar el dataset de cuadros etiquetados en BORIS y entrenar ResNet-18; medir kappa, F1 y MAE contra el gold standard; y entregar el sistema dockerizado con manuales de usuario y el material completo de defensa.")

# ── DIAPOSITIVA 21 — Bibliografía ────────────────────────────
add_slide_heading(doc, 21, "Bibliografía")
add_tiempo(doc, 35)
add_note(doc, "Transición breve — no se leen todas las referencias en voz alta.")
A(doc, "La bibliografía incluye las doce fuentes más relevantes: el artículo fundacional de Porsolt de 1977, ByteTrack del ECCV 2022, las redes residuales de He y colaboradores, YOLOv8 de Ultralytics, CLAHE de Zuiderveld, las normas ISO/IEC 25010 e ISO/IEC/IEEE 12207, BORIS de Friard y Gamba, ezTrack de Pennington, el reporte de depresión de la OMS, la Guía Scrum y la NOM-062-ZOO-1999. La bibliografía completa se incluye en el documento escrito del Trabajo Terminal.")

# ── DIAPOSITIVA 22 — Cierre ───────────────────────────────────
add_slide_heading(doc, 22, "Cierre")
add_tiempo(doc, 30)
add_note(doc, "Ambos se colocan juntos frente al estrado.")
V(doc, "Para cerrar: nuestro objetivo es claro —pasar de horas frente al cronómetro a minutos de análisis automatizado. Un prototipo abierto, desplegable con Docker y validado contra el gold standard del laboratorio.")
A(doc, "Muchas gracias por su atención. Quedamos a su disposición para responder sus preguntas.")
add_note(doc, "Ambos aguardan en silencio a que el jurado formule sus preguntas.")

# ══════════════════════════════════════════════════════════════
# SECCIÓN: Preguntas frecuentes
# ══════════════════════════════════════════════════════════════
doc.add_page_break()

p = doc.add_paragraph()
para_spacing(p, before=8, after=6)
set_para_border(p, 'bottom', '0D1B3E', size=6, space=4)
run_add(p, "Posibles preguntas del jurado y respuestas de apoyo", bold=True, color=NAVY, size=16)

faq = [
    ("¿Por qué usar ByteTrack y no DeepSORT?",
     "ByteTrack evita descartar las detecciones de baja confianza: las asocia en un segundo paso con las trayectorias existentes, lo que reduce pérdidas de identidad durante oclusiones. En el FST, donde las ratas pueden solaparse parcialmente, esto es una ventaja directa sobre DeepSORT, que simplemente descarta esas detecciones."),
    ("¿Por qué ResNet-18 y no una arquitectura más moderna como ViT?",
     "El sistema debe operar en la CPU del equipo del laboratorio sin GPU. ResNet-18 ofrece el mejor balance precisión-velocidad para ese requisito. Los Vision Transformers tienen una huella computacional significativamente mayor, inviable sin GPU. Si el hardware lo permite, ResNet-50 está disponible como alternativa."),
    ("¿Por qué la kappa de Cohen como métrica de concordancia?",
     "Es la métrica estándar en estudios conductuales para medir la concordancia inter-evaluador corrigiendo el acuerdo por azar. La complementamos con el F1-score por clase —para detectar sesgos en conductas minoritarias— y el MAE en segundos —para la precisión temporal—. Esta combinación cubre tanto la concordancia categórica como la cuantificación temporal."),
    ("¿Qué ocurre si el video no cumple las condiciones de grabación?",
     "El sistema intenta corregir el contraste con CLAHE antes de la detección. Si la confianza en la localización de los cilindros sigue por debajo de 0.70, el pipeline genera un reporte de diagnóstico en PDF indicando el cuadro problemático y la etapa de falla, conforme a RN-11. El investigador siempre recibe retroalimentación accionable en lugar de un fallo silencioso."),
    ("¿Por qué JWT y no sesiones de servidor?",
     "JWT permite una arquitectura sin estado en el servidor, facilita el despliegue con Docker Compose y la escalabilidad futura. Dado que el frontend es una SPA en React separada del backend Flask, las sesiones de servidor añadirían complejidad de CORS y cookies de sesión innecesaria para el prototipo."),
    ("¿Cómo garantizan que los resultados sean reproducibles entre experimentos?",
     "Las reglas de negocio formalizan el protocolo: un solo video por slot, clasificación mutuamente excluyente cuadro a cuadro y validación contra el gold standard con kappa. El mismo pipeline y los mismos parámetros se aplican a todos los videos. Los resultados se conservan indefinidamente (RN-06) para permitir comparaciones históricas."),
]

for q, r in faq:
    p_q = doc.add_paragraph()
    para_spacing(p_q, before=8, after=3)
    run_add(p_q, f"P: {q}", bold=True, color=GREEN, size=11)
    p_r = doc.add_paragraph()
    para_spacing(p_r, before=0, after=8)
    run_add(p_r, f"R: {r}", color=GRAY, size=11)

# ── Guardar ───────────────────────────────────────────────────
doc.save(OUT)
print("OK:", OUT)
